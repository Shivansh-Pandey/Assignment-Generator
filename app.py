import os
import io
import json
import re
import base64
from datetime import datetime
from flask import Flask, render_template, request, jsonify
from google import genai
from google.genai import types
from dotenv import load_dotenv

# ReportLab
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, KeepTogether
)
from reportlab.lib.colors import black, white
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ─────────────────────────────────────────────
# Register TrueType Fonts for Native Unicode Support (μ, π, ×, α, β, √, etc.)
# ─────────────────────────────────────────────
_TIMES_FONT = "Times-Roman"
_TIMES_BOLD = "Times-Bold"
_TIMES_ITALIC = "Times-Italic"
_TIMES_BOLD_ITALIC = "Times-BoldItalic"

_win_fonts = r"C:\Windows\Fonts"
if os.path.exists(os.path.join(_win_fonts, "times.ttf")):
    try:
        pdfmetrics.registerFont(TTFont("CustomTimes", os.path.join(_win_fonts, "times.ttf")))
        pdfmetrics.registerFont(TTFont("CustomTimes-Bold", os.path.join(_win_fonts, "timesbd.ttf")))
        pdfmetrics.registerFont(TTFont("CustomTimes-Italic", os.path.join(_win_fonts, "timesi.ttf")))
        pdfmetrics.registerFont(TTFont("CustomTimes-BoldItalic", os.path.join(_win_fonts, "timesbi.ttf")))
        _TIMES_FONT = "CustomTimes"
        _TIMES_BOLD = "CustomTimes-Bold"
        _TIMES_ITALIC = "CustomTimes-Italic"
        _TIMES_BOLD_ITALIC = "CustomTimes-BoldItalic"
    except Exception:
        pass

# python-docx
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

load_dotenv()
app = Flask(__name__)

# ─────────────────────────────────────────────
# Gemini client
# ─────────────────────────────────────────────
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
client = genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None

# ─────────────────────────────────────────────
# PDF layout constants
# ─────────────────────────────────────────────
PAGE_W, PAGE_H = A4
MARGIN_L = 20 * mm
MARGIN_R = 20 * mm
MARGIN_T = 32 * mm   # 3-line header on page 1
MARGIN_B = 22 * mm

SCHOOL_NAME = "Eicher School Faridabad"
SCHOOL_CODE = "ESF"
LOGO_PATH   = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Logo.png")

_ROMAN = {
    "Class 1":"I",  "Class 2":"II",  "Class 3":"III",
    "Class 4":"IV", "Class 5":"V",   "Class 6":"VI",
    "Class 7":"VII","Class 8":"VIII","Class 9":"IX",
    "Class 10":"X", "Class 11":"XI", "Class 12":"XII",
}


def _clean_topic(raw_title: str, topic_fallback: str = "") -> str:
    """Strips any level indicators, assignment prefixes, class names, etc.
    Returns clean topic string."""
    text = raw_title.strip()
    text = re.sub(r'\s*\([^)]*level[^)]*\)', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\([^)]*jee[^)]*\)', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\([^)]*olympiad[^)]*\)', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\([^)]*cbse[^)]*\)', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\([^)]*board[^)]*\)', '', text, flags=re.IGNORECASE)

    text = re.sub(r'^(?:Class\s+\d+\s+)?[A-Za-z\s]+Assignment:\s*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'^(?:Advanced|Standard|Olympiad)\s+Assignment:\s*', '', text, flags=re.IGNORECASE)
    text = text.strip()
    return text if text else topic_fallback.strip()

_GREEK_UNICODE_MAP = {
    # Alpha
    'α': '<font name="Symbol">a</font>', '\u03b1': '<font name="Symbol">a</font>', '\u0251': '<font name="Symbol">a</font>', 'Α': '<font name="Symbol">A</font>',
    # Beta
    'β': '<font name="Symbol">b</font>', '\u03b2': '<font name="Symbol">b</font>', 'Β': '<font name="Symbol">B</font>',
    # Gamma
    'γ': '<font name="Symbol">g</font>', '\u03b3': '<font name="Symbol">g</font>', 'Γ': '<font name="Symbol">G</font>',
    # Delta
    'δ': '<font name="Symbol">d</font>', '\u03b4': '<font name="Symbol">d</font>', 'Δ': '<font name="Symbol">D</font>', '\u0394': '<font name="Symbol">D</font>',
    # Epsilon
    'ε': '<font name="Symbol">e</font>', '\u03b5': '<font name="Symbol">e</font>', '\u03f5': '<font name="Symbol">e</font>', 'Ε': '<font name="Symbol">E</font>',
    # Zeta
    'ζ': '<font name="Symbol">z</font>', '\u03b6': '<font name="Symbol">z</font>', 'Ζ': '<font name="Symbol">Z</font>',
    # Eta
    'η': '<font name="Symbol">h</font>', '\u03b7': '<font name="Symbol">h</font>', 'Η': '<font name="Symbol">H</font>',
    # Theta
    'θ': '<font name="Symbol">q</font>', '\u03b8': '<font name="Symbol">q</font>', '\u03d1': '<font name="Symbol">q</font>', 'Θ': '<font name="Symbol">Q</font>',
    # Iota
    'ι': '<font name="Symbol">i</font>', '\u03b9': '<font name="Symbol">i</font>', 'Ι': '<font name="Symbol">I</font>',
    # Kappa
    'κ': '<font name="Symbol">k</font>', '\u03ba': '<font name="Symbol">k</font>', 'Κ': '<font name="Symbol">K</font>',
    # Lambda
    'λ': '<font name="Symbol">l</font>', '\u03bb': '<font name="Symbol">l</font>', 'Λ': '<font name="Symbol">L</font>',
    # Mu (BOTH \u03bc Greek mu AND \u00b5 Micro sign!)
    'μ': '<font name="Symbol">m</font>', '\u03bc': '<font name="Symbol">m</font>', 'µ': '<font name="Symbol">m</font>', '\u00b5': '<font name="Symbol">m</font>', 'Μ': '<font name="Symbol">M</font>',
    # Nu
    'ν': '<font name="Symbol">n</font>', '\u03bd': '<font name="Symbol">n</font>', 'Ν': '<font name="Symbol">N</font>',
    # Xi
    'ξ': '<font name="Symbol">x</font>', '\u03be': '<font name="Symbol">x</font>', 'Ξ': '<font name="Symbol">X</font>',
    # Pi
    'π': '<font name="Symbol">p</font>', '\u03c0': '<font name="Symbol">p</font>', 'Π': '<font name="Symbol">P</font>',
    # Rho
    'ρ': '<font name="Symbol">r</font>', '\u03c1': '<font name="Symbol">r</font>', 'Ρ': '<font name="Symbol">P</font>',
    # Sigma
    'σ': '<font name="Symbol">s</font>', '\u03c3': '<font name="Symbol">s</font>', '\u03c2': '<font name="Symbol">s</font>', 'Σ': '<font name="Symbol">S</font>',
    # Tau
    'τ': '<font name="Symbol">t</font>', '\u03c4': '<font name="Symbol">t</font>', 'Τ': '<font name="Symbol">T</font>',
    # Upsilon
    'υ': '<font name="Symbol">u</font>', '\u03c5': '<font name="Symbol">u</font>', 'Υ': '<font name="Symbol">U</font>',
    # Phi
    'φ': '<font name="Symbol">f</font>', '\u03c6': '<font name="Symbol">f</font>', 'ϕ': '<font name="Symbol">f</font>', '\u03d5': '<font name="Symbol">f</font>', 'Φ': '<font name="Symbol">F</font>',
    # Chi
    'χ': '<font name="Symbol">c</font>', '\u03c7': '<font name="Symbol">c</font>', 'Χ': '<font name="Symbol">C</font>',
    # Psi
    'ψ': '<font name="Symbol">y</font>', '\u03c8': '<font name="Symbol">y</font>', 'Ψ': '<font name="Symbol">Y</font>',
    # Omega
    'ω': '<font name="Symbol">w</font>', '\u03c9': '<font name="Symbol">w</font>', 'Ω': '<font name="Symbol">W</font>', '\u03a9': '<font name="Symbol">W</font>',
}

_SUB_SUPER_UNICODE_MAP = {
    '₀': '<sub>0</sub>', '₁': '<sub>1</sub>', '₂': '<sub>2</sub>',
    '₃': '<sub>3</sub>', '₄': '<sub>4</sub>', '₅': '<sub>5</sub>',
    '₆': '<sub>6</sub>', '₇': '<sub>7</sub>', '₈': '<sub>8</sub>',
    '₉': '<sub>9</sub>', 'ᵢ': '<sub>i</sub>', 'ₓ': '<sub>x</sub>',
    'ᵧ': '<sub>y</sub>',
    '⁰': '<sup>0</sup>', '¹': '<sup>1</sup>', '²': '<sup>2</sup>',
    '³': '<sup>3</sup>', '⁴': '<sup>4</sup>', '⁵': '<sup>5</sup>',
    '⁶': '<sup>6</sup>', '⁷': '<sup>7</sup>', '⁸': '<sup>8</sup>',
    '⁹': '<sup>9</sup>', 'ⁿ': '<sup>n</sup>', '⁻': '<sup>-</sup>',
    '⁺': '<sup>+</sup>',
}


def _clean_text_noise(text: str) -> str:
    """Strips leading dots, bullet noise, and black box characters."""
    text = re.sub(r'^\s*[\.\•\-\*]\s*', '', text)
    for b in ['■', '□', '▪', '▫', '●', '○', '◆', '◇', '￼', '']:
        text = text.replace(b, '')
    return text.strip()


def _convert_unicode_symbols(text: str) -> str:
    """Converts Greek letters, subscripts, superscripts, and vector arrows to native Unicode chars."""
    for ch, rep in _SUB_SUPER_UNICODE_MAP.items():
        text = text.replace(ch, rep)
    text = re.sub(r'([A-Za-z])\u20d7', r'<b><i>\1</i></b>', text)
    return text


def _format_vectors(text: str) -> str:
    """Formats vector notation like vec(B), \\vec{B}, vec B to NCERT Bold-Italic (<b><i>B</i></b>)."""
    text = re.sub(r'\\vec\{([A-Za-z0-9]+)\}', r'<b><i>\1</i></b>', text)
    text = re.sub(r'\\vec\s+([A-Za-z0-9])', r'<b><i>\1</i></b>', text)
    text = re.sub(r'\bvec\{([A-Za-z0-9]+)\}', r'<b><i>\1</i></b>', text)
    text = re.sub(r'\bvec\(([^)]+)\)', r'<b><i>\1</i></b>', text)
    text = re.sub(r'\bvec\s+([A-Za-z0-9])\b', r'<b><i>\1</i></b>', text)
    return text


def _format_subparts(text: str) -> str:
    """Inserts line breaks and non-breaking space indents (\u00a0) for sub-parts."""
    indent4 = '\u00a0\u00a0\u00a0\u00a0'
    indent8 = '\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0'
    text = re.sub(r'(\s*(?:\[|\()?[\d\s]*marks?(?:\]|\))?\.?\s*)(\([a-e]\))\s+', r'<br/>' + indent4 + r'\2 ', text, flags=re.IGNORECASE)
    text = re.sub(r'(\s+)(\((?:i|ii|iii|iv|v|vi)\))\s+', r'<br/>' + indent8 + r'\1 ', text, flags=re.IGNORECASE)
    return text


def _format_physics_math(text: str) -> str:
    """Converts mu0, pi, theta, subscripts and superscripts to clean mathematical notation."""
    # 1. mu0 / mu_0 / \mu0 / \mu_0 / μ0
    text = re.sub(r'\\?mu_?0|&mu;_?0|μ0', 'μ₀', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?mu\b|&mu;', 'μ', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?pi\b|&pi;', 'π', text, flags=re.IGNORECASE)

    # 2. theta / \theta with or without subscripts (e.g. theta_1, theta_2, theta_a)
    text = re.sub(r'\\?theta_?([0-9a-zA-Z]+)\b', r'θ<sub>\1</sub>', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?theta\b|&theta;', 'θ', text, flags=re.IGNORECASE)

    # 3. Other Greek symbols with optional subscripts
    text = re.sub(r'\\?phi_?([0-9a-zA-Z]+)\b', r'φ<sub>\1</sub>', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?phi\b', 'φ', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?Phi\b', 'Φ', text)

    text = re.sub(r'\\?alpha_?([0-9a-zA-Z]+)\b', r'α<sub>\1</sub>', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?alpha\b|&alpha;', 'α', text, flags=re.IGNORECASE)

    text = re.sub(r'\\?beta_?([0-9a-zA-Z]+)\b', r'β<sub>\1</sub>', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?beta\b|&beta;', 'β', text, flags=re.IGNORECASE)

    text = re.sub(r'\\?gamma_?([0-9a-zA-Z]+)\b', r'γ<sub>\1</sub>', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?gamma\b|&gamma;', 'γ', text, flags=re.IGNORECASE)

    text = re.sub(r'\\?omega\b|&omega;', 'ω', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?Omega\b|&Omega;|ohms?\b', 'Ω', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?Delta\b|&Delta;', 'Δ', text)
    text = re.sub(r'\\?delta\b|&delta;', 'δ', text)
    text = re.sub(r'\\?lambda\b|&lambda;', 'λ', text)
    text = re.sub(r'\\?rho\b|&rho;', 'ρ', text)
    text = re.sub(r'\\?sigma\b|&sigma;', 'σ', text)
    text = re.sub(r'\\?epsilon_?0|\\?eps_?0', 'ε₀', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?epsilon\b|\\?eps\b', 'ε', text, flags=re.IGNORECASE)
    text = re.sub(r'\\?tau\b', 'τ', text)

    # General Subscripts & Superscripts
    text = re.sub(r'\b([A-Za-z]+)_([0-9a-zA-Z]+)\b', r'\1<sub>\2</sub>', text)
    text = re.sub(r'\b([A-Za-z])\^?([0-9])\b', r'\1<sup>\2</sup>', text)

    # Square root
    text = re.sub(r'sqrt\s*\(([^)]+)\)', r'√(\1)', text)
    text = re.sub(r'sqrt\s*([0-9a-zA-Z]+)', r'√\1', text)
    return text


def _format_multiplication(text: str) -> str:
    """Replaces '*' with clean physics multiplication sign '×'."""
    text = re.sub(r'(\S+)\s*\*\s*(\S+)', r'\1 × \2', text)
    text = text.replace(' * ', ' × ')
    return text


def _smart_xml_escape(text: str) -> str:
    """Escapes & < > except for valid ReportLab HTML tags (<sub>, <super>, <sup>, <b>, <i>, <br/>)."""
    tags = []
    def _save_tag(m):
        tags.append(m.group(0))
        return f"___TAG_{len(tags)-1}___"

    pattern = r'</?(?:sub|super|sup|b|i|br\s*/?)>'
    text = re.sub(pattern, _save_tag, text, flags=re.IGNORECASE)

    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')

    for i, tag in enumerate(tags):
        text = text.replace(f"___TAG_{i}___", tag)

    return text


def _sanitize(text: str) -> str:
    """Pass-through for TrueType Unicode support."""
    return text


def _prep(text: str) -> str:
    """Full preparation pipeline for questions, instructions, and choices."""
    text = _clean_text_noise(text)
    text = _convert_unicode_symbols(text)
    text = _format_vectors(text)
    text = _format_subparts(text)
    text = _format_physics_math(text)
    text = _format_multiplication(text)
    text = _smart_xml_escape(text)
    return text


def _prep_mcq_options(raw_options: list) -> list:
    """Ensures MCQ options have bold, clean A., B., C., D. labels."""
    labels = ["A.", "B.", "C.", "D."]
    out = []
    for idx, opt in enumerate(raw_options):
        cleaned = _prep(opt)
        cleaned = re.sub(r'^(?:&lt;b&gt;)?\s*[A-D][\.\)]\s*(?:&lt;/b&gt;)?\s*', '', cleaned, flags=re.IGNORECASE)
        prefix = labels[idx] if idx < len(labels) else f"{chr(65+idx)}."
        out.append(f"<b>{prefix}</b> {cleaned}")
    return out


# ─────────────────────────────────────────────
# Difficulty descriptors
# ─────────────────────────────────────────────
_LEVEL_DESCRIPTORS = {
    "standard": (
        "STANDARD (NCERT Baseline)",
        """- Questions must be directly based on the NCERT textbook for this class and chapter.
- Use simple, clear language appropriate for a regular school exam.
- MCQs should test direct recall or one-step reasoning from NCERT content.
- Short/medium/long answers should match NCERT exercise questions in difficulty.
- Do NOT use topics outside the NCERT syllabus for this class."""
    ),
    "advanced": (
        "ADVANCED (CBSE Board / Competitive Prep)",
        """- Questions should go beyond simple recall; include application and analysis-level questions.
- Style and difficulty should match CBSE Board Exam Previous Year Questions (PYQs).
- MCQs may involve case-based or assertion-reason format as per latest CBSE pattern.
- Medium and long answers should require multi-step reasoning and structured responses.
- Stay within the CBSE syllabus but push to the harder end of board-level difficulty."""
    ),
    "olympiad": (
        "OLYMPIAD / JEE LEVEL (High-Order Thinking)",
        """- Questions must be at the difficulty level of JEE Main, JEE Advanced, NTSE, or Science Olympiad PYQs.
- MCQs may have more than one correct option, or involve multi-concept integration.
- Problems should require deep conceptual understanding, multi-step derivations, or novel application.
- For science/math topics: include numerical problems similar to JEE Main / JEE Advanced style.
- The content must still be rooted in the CBSE/NCERT syllabus for the given class."""
    ),
}

# ─────────────────────────────────────────────
# Prompt builder
# ─────────────────────────────────────────────
def build_prompt(class_level, topic, marks_2, marks_3, marks_5, mcq,
                 difficulty_level="standard", subtopics="",
                 exclusions="", custom_instructions=""):

    level_key = difficulty_level.lower() if difficulty_level.lower() in _LEVEL_DESCRIPTORS else "standard"
    level_label, level_instr = _LEVEL_DESCRIPTORS[level_key]

    sub_block = f"""
===================================================
SUB-TOPICS TO FOCUS ON  <- PRIORITISE THESE
===================================================
{subtopics.strip()}
Every listed sub-topic must appear in at least one question.""" if subtopics.strip() else ""

    excl_block = f"""
===================================================
TOPICS / FORMATS TO EXCLUDE  <- DO NOT INCLUDE
===================================================
{exclusions.strip()}
Zero questions related to any excluded item.""" if exclusions.strip() else ""

    cust_block = f"""
===================================================
CUSTOM INSTRUCTIONS  <- OVERRIDE DEFAULTS IF NEEDED
===================================================
{custom_instructions.strip()}""" if custom_instructions.strip() else ""

    return f"""
You are an expert Indian school examiner with deep knowledge of the NCERT and CBSE syllabus.
Generate a formal school assignment strictly following the specification below.

===================================================
ASSIGNMENT SPECIFICATION
===================================================
- Class / Grade Level : {class_level}   <- STRICTLY follow this class level
- Topic               : {topic}
- Difficulty Level    : {level_label}
- MCQ questions       : {mcq}  (1 mark each)
- 2-mark questions    : {marks_2}
- 3-mark questions    : {marks_3}
- 5-mark questions    : {marks_5}
{sub_block}{excl_block}{cust_block}

===================================================
CURRICULUM & DIFFICULTY RULES  <- MANDATORY
===================================================
1. ALL questions must be grounded in the official NCERT/CBSE syllabus for {class_level}.
2. Do NOT include concepts or topics that belong to a different class or grade.
3. Difficulty calibration:
{level_instr}

===================================================
FORMATTING & PRESENTATION RULES  <- MANDATORY
===================================================
1. QUESTION SUB-PARTS (CRITICAL):
   - When a question has sub-parts, format each sub-part with a clear label starting on a NEW line:
     (a) ...
     (b) ...
     (c) ...
   - Sub-items like (i), (ii), (iii) MUST also be formatted on new lines.
   - Example format for multi-part questions:
     "(a) State Ampere's Circuital Law. (1 mark)\\n(b) Derive the expression for the magnetic field inside a solenoid. (2 marks)\\n(c) A coaxial cable carries current I. Find the magnetic field for:\\n(i) r < a\\n(ii) a < r < b\\n(iii) r > b"

2. SCIENTIFIC & MATHEMATICAL NOTATION:
   - Write clear mathematical formulas and expressions.
   - Use subscripts like mu_0, B_0, I_1, r_2, v_d and superscripts like N^2, r^2, 10^-6.
   - Use standard names for Greek letters (mu_0, pi, theta, omega, lambda, delta, ohm).
   - Use clean fraction notation like "a / b" or "(mu_0 * I * a) / (2 * pi)".
   - NEVER output broken text or truncated option letters.

3. MCQ OPTIONS:
   - Provide exactly 4 options per MCQ.
   - Start each option text with "A. ", "B. ", "C. ", "D. ".

===================================================
OUTPUT FORMAT
===================================================
Return ONLY a valid JSON object (no markdown fences, no extra text):

{{
  "title": "<Clean chapter/topic name ONLY - do NOT include class name, 'Assignment' prefix, or difficulty level string in title>",
  "subject": "<NCERT subject name e.g. Physics, Chemistry, Mathematics, Biology, Science>",
  "class_level": "{class_level}",
  "topic": "{topic}",
  "difficulty": "{level_label}",
  "total_marks": <integer>,
  "sections": [
    {{
      "section_label": "Section A - Multiple Choice Questions",
      "instructions": "Choose the correct option. Each question carries 1 mark.",
      "questions": [
        {{
          "number": 1, "marks": 1, "type": "MCQ",
          "question": "<question text>",
          "options": ["A. ...", "B. ...", "C. ...", "D. ..."],
          "answer": "A"
        }}
      ]
    }},
    {{
      "section_label": "Section B - Short Answer Questions",
      "instructions": "Answer in 2-3 sentences. Each question carries 2 marks.",
      "questions": [
        {{"number": <n>, "marks": 2, "type": "SHORT", "question": "<text>", "options": [], "answer": ""}}
      ]
    }},
    {{
      "section_label": "Section C - Medium Answer Questions",
      "instructions": "Answer in 3-5 sentences. Each question carries 3 marks.",
      "questions": [
        {{"number": <n>, "marks": 3, "type": "MEDIUM", "question": "<text>", "options": [], "answer": ""}}
      ]
    }},
    {{
      "section_label": "Section D - Long Answer Questions",
      "instructions": "Answer in detail. Each question carries 5 marks.",
      "questions": [
        {{"number": <n>, "marks": 5, "type": "LONG", "question": "<text>", "options": [], "answer": ""}}
      ]
    }}
  ],
  "marking_scheme": [
    {{
      "question_number": <n>, "marks": <int>, "type": "<MCQ|SHORT|MEDIUM|LONG>",
      "question_summary": "<brief>",
      "answer_key": "<answer or correct option>",
      "mark_breakdown": ["<step 1 - X mark(s)>", "<step 2 - Y mark(s)>"]
    }}
  ]
}}

Mandatory rules:
- Include exactly {mcq} MCQ questions (omit Section A if 0).
- Include exactly {marks_2} 2-mark short questions (omit Section B if 0).
- Include exactly {marks_3} 3-mark questions (omit Section C if 0).
- Include exactly {marks_5} 5-mark questions (omit Section D if 0).
- Only include sections with at least one question.
- Number questions sequentially across ALL sections (1, 2, 3 ...).
- marking_scheme must cover EVERY question.
- JSON must be syntactically valid.
- All content must be from {class_level} NCERT/CBSE syllabus for topic "{topic}".
""".strip()


def _clean_json_string(raw: str) -> str:
    """Strips markdown fences and extracts JSON object."""
    raw = re.sub(r'^```(?:json)?\s*', '', raw, flags=re.MULTILINE)
    raw = re.sub(r'\s*```$', '', raw, flags=re.MULTILINE)
    raw = raw.strip()
    match = re.search(r'(\{[\s\S]*\})', raw)
    if match:
        raw = match.group(1)
    return raw


# ─────────────────────────────────────────────
# Gemini
# ─────────────────────────────────────────────
def call_gemini(prompt: str) -> dict:
    if not client:
        raise ValueError("GEMINI_API_KEY_MISSING: API key not set in .env file.")

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.4,
                max_output_tokens=16384,
                response_mime_type="application/json",
            ),
        )
        raw = response.text.strip() if response and response.text else ""
        if raw:
            cleaned = _clean_json_string(raw)
            try:
                return json.loads(cleaned)
            except Exception:
                repaired = re.sub(r',\s*([\}\]])', r'\1', cleaned)
                return json.loads(repaired, strict=False)
    except Exception as e:
        print(f"[Gemini Call 1 Warning] {e}")

    # Fallback call without strict response_mime_type constraint if first call failed
    response2 = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.3,
            max_output_tokens=16384,
        ),
    )
    if not response2 or not response2.text:
        raise ValueError("GEMINI_EMPTY_RESPONSE: The AI returned an empty response. Try simplifying the topic or reducing question counts.")

    cleaned2 = _clean_json_string(response2.text.strip())
    try:
        return json.loads(cleaned2)
    except Exception:
        repaired2 = re.sub(r',\s*([\}\]])', r'\1', cleaned2)
        return json.loads(repaired2, strict=False)




# ─────────────────────────────────────────────
# PDF styles
# ─────────────────────────────────────────────
def _get_styles():
    title_sub_style = ParagraphStyle(
        "AssignTitleSub", fontSize=12, leading=16, alignment=TA_CENTER,
        fontName=_TIMES_BOLD, textColor=black, spaceAfter=2,
    )
    title_style = ParagraphStyle(
        "AssignTitle", fontSize=14, leading=18, alignment=TA_CENTER,
        fontName=_TIMES_BOLD, textColor=black, spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        "MetaLine", fontSize=9, leading=12, alignment=TA_CENTER,
        fontName=_TIMES_FONT, textColor=black,
    )
    section_style = ParagraphStyle(
        "SectionLabel", fontSize=11, leading=14, alignment=TA_LEFT,
        fontName=_TIMES_BOLD, textColor=black, spaceBefore=8, spaceAfter=2,
    )
    instructions_style = ParagraphStyle(
        "Instructions", fontSize=9, leading=12, alignment=TA_LEFT,
        fontName=_TIMES_ITALIC, textColor=black, spaceAfter=4,
    )
    question_style = ParagraphStyle(
        "Question", fontSize=10.5, leading=15, alignment=TA_JUSTIFY,
        fontName=_TIMES_FONT, textColor=black,
    )
    option_style = ParagraphStyle(
        "Option", fontSize=10, leading=13, alignment=TA_LEFT,
        fontName=_TIMES_FONT, textColor=black, leftIndent=18,
    )
    ms_head_style = ParagraphStyle(
        "MSHead", fontSize=12, leading=15, alignment=TA_CENTER,
        fontName=_TIMES_BOLD, textColor=black, spaceBefore=4, spaceAfter=6,
    )
    ms_q_style = ParagraphStyle(
        "MSQuestion", fontSize=9, leading=12, alignment=TA_LEFT,
        fontName=_TIMES_BOLD, textColor=black, spaceBefore=4,
    )
    ms_answer_style = ParagraphStyle(
        "MSAnswer", fontSize=9, leading=12, alignment=TA_LEFT,
        fontName=_TIMES_FONT, textColor=black, leftIndent=10,
    )
    ms_step_style = ParagraphStyle(
        "MSStep", fontSize=9, leading=12, alignment=TA_LEFT,
        fontName=_TIMES_FONT, textColor=black, leftIndent=18,
    )
    return {
        "title": title_style, "title_sub": title_sub_style, "meta": meta_style,
        "section": section_style, "instructions": instructions_style,
        "question": question_style, "option": option_style,
        "ms_head": ms_head_style, "ms_q": ms_q_style,
        "ms_answer": ms_answer_style, "ms_step": ms_step_style,
    }


# ─────────────────────────────────────────────
# Page callbacks factory
# Returns (first_page_cb, later_pages_cb)
# Header ONLY on page 1, footer on ALL pages
# ─────────────────────────────────────────────
def make_callbacks(subject: str, class_level: str, topic: str,
                   month_year: str, total_pages: dict):
    roman    = _ROMAN.get(class_level, class_level.upper().replace("CLASS ", ""))
    subcode  = re.sub(r'\s+', '', subject.upper())[:12]
    now      = datetime.now()
    ay_start = now.year if now.month >= 4 else now.year - 1
    acad_yr  = f"{ay_start}-{str(ay_start + 1)[-2:]}"
    ref_code = f"{SCHOOL_CODE}/ASB/{subcode}/{roman}/{now.strftime('%B').upper()} {acad_yr}"

    def _draw_footer(canvas, doc):
        pw = doc.pagesize[0]
        fy = MARGIN_B - 3 * mm
        canvas.setStrokeColor(black)
        canvas.setLineWidth(0.7)
        canvas.line(MARGIN_L, fy, pw - MARGIN_R, fy)
        canvas.setFont(_TIMES_FONT, 7.5)
        canvas.drawString(MARGIN_L, fy - 4 * mm, ref_code)
        tp      = total_pages.get("n", 0)
        pg_text = f"Page {doc.page} of {tp}" if tp else f"Page {doc.page}"
        canvas.drawRightString(pw - MARGIN_R, fy - 4 * mm, pg_text)

    def _first_page(canvas, doc):
        canvas.saveState()
        pw = doc.pagesize[0]
        ph = doc.pagesize[1]

        # ── HEADER (page 1 only) ──────────────────────────────
        LOGO_W     = 20 * mm
        LOGO_H     = 20 * mm
        header_top = ph - 4 * mm
        header_bot = ph - MARGIN_T + 4 * mm
        header_h   = header_top - header_bot

        canvas.setStrokeColor(black)
        canvas.setLineWidth(0.8)
        canvas.rect(MARGIN_L, header_bot,
                    pw - MARGIN_L - MARGIN_R, header_h, stroke=1, fill=0)

        # Logo – left, vertically centred
        logo_y = header_bot + (header_h - LOGO_H) / 2
        if os.path.exists(LOGO_PATH):
            canvas.drawImage(LOGO_PATH, MARGIN_L + 2 * mm, logo_y,
                             width=LOGO_W, height=LOGO_H,
                             preserveAspectRatio=True, mask="auto")
        else:
            canvas.setFont("Helvetica-Bold", 7)
            canvas.drawCentredString(MARGIN_L + 2*mm + LOGO_W/2,
                                     logo_y + LOGO_H/2 - 3, "ESF")

        # Text – centred in remaining width
        txt_x = MARGIN_L + LOGO_W + 7 * mm
        cx    = txt_x + (pw - txt_x - MARGIN_R - 3 * mm) / 2

        # Vertical centering for 3 rows
        total_text_h = 7 + 5.5 + 5
        start_y = header_bot + header_h / 2 + (total_text_h / 2) * mm

        r1y = start_y
        canvas.setFont(_TIMES_BOLD, 13)
        canvas.drawCentredString(cx, r1y, SCHOOL_NAME)

        r2y = r1y - 6 * mm
        canvas.setFont(_TIMES_BOLD, 11)
        canvas.drawCentredString(cx, r2y, f"{subject} Assignment")

        r3y = r2y - 5.5 * mm
        canvas.setFont("Times-Roman", 10)
        canvas.drawCentredString(cx, r3y, f"{class_level}   |   {month_year}")

        _draw_footer(canvas, doc)
        canvas.restoreState()

    def _later_pages(canvas, doc):
        canvas.saveState()
        _draw_footer(canvas, doc)
        canvas.restoreState()

    return _first_page, _later_pages


# ─────────────────────────────────────────────
# Story builder
# ─────────────────────────────────────────────
def _build_story(data: dict, styles: dict) -> list:
    story = []
    W = PAGE_W - MARGIN_L - MARGIN_R
    total_marks = data.get("total_marks", "-")
    subject     = data.get("subject", "Physics")

    raw_title   = data.get("title", "")
    topic_input = data.get("topic", "")
    topic_clean = _clean_topic(raw_title, topic_input)

    # Document Header Title (Physics Assignment / Clean Topic Name)
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph(f"<b>{_prep(subject)} Assignment</b>", styles["title_sub"]))
    story.append(Paragraph(f"<b>{_prep(topic_clean)}</b>", styles["title"]))
    story.append(Spacer(1, 3 * mm))
    story.append(HRFlowable(width=W, thickness=0.8, color=black))


    # Sections & Questions
    for section in data.get("sections", []):
        story.append(Spacer(1, 3 * mm))
        story.append(Paragraph(_sanitize(section.get("section_label", "")), styles["section"]))
        story.append(HRFlowable(width=W, thickness=0.6, color=black))
        story.append(Paragraph(_sanitize(section.get("instructions", "")), styles["instructions"]))

        for q in section.get("questions", []):
            marks  = q.get("marks", 1)
            q_text = _prep(q.get("question", ""))
            q_type = q.get("type", "")
            q_num  = q.get("number", "")

            q_row = [
                Paragraph(f"<b>Q{q_num}.</b> {q_text}", styles["question"]),
                Paragraph(
                    f"[{marks} {'Mark' if marks == 1 else 'Marks'}]",
                    ParagraphStyle("MarksR", fontSize=9, alignment=TA_RIGHT,
                                   fontName="Times-Bold", textColor=black)
                ),
            ]
            q_table = Table([q_row], colWidths=[W - 22*mm, 22*mm])
            q_table.setStyle(TableStyle([
                ("VALIGN",        (0,0),(-1,-1), "TOP"),
                ("LEFTPADDING",   (0,0),(-1,-1), 0),
                ("RIGHTPADDING",  (0,0),(-1,-1), 0),
                ("TOPPADDING",    (0,0),(-1,-1), 3),
                ("BOTTOMPADDING", (0,0),(-1,-1), 2),
            ]))

            items = [q_table]

            if q_type == "MCQ":
                options = _prep_mcq_options(q.get("options", []))
                if options:
                    # Strip tags to check length of raw text
                    raw_lens = [len(re.sub(r'<[^>]+>', '', o)) for o in options]
                    is_long  = any(l > 26 for l in raw_lens) or any('\n' in o or '<br' in o for o in options)

                    if is_long:
                        # Stacked 1-column layout for math expressions / long options
                        for opt in options:
                            ot = Table([[Paragraph(opt, styles["option"])]], colWidths=[W])
                            ot.setStyle(TableStyle([
                                ("LEFTPADDING",   (0,0),(-1,-1), 18),
                                ("TOPPADDING",    (0,0),(-1,-1), 1.5),
                                ("BOTTOMPADDING", (0,0),(-1,-1), 1.5),
                            ]))
                            items.append(ot)
                    else:
                        # 2-column side-by-side layout for short options
                        opt_pairs = [options[i:i+2] for i in range(0, len(options), 2)]
                        for pair in opt_pairs:
                            row = [Paragraph(o, styles["option"]) for o in pair]
                            if len(row) == 1:
                                row.append(Paragraph("", styles["option"]))
                            ot = Table([row], colWidths=[W/2, W/2])
                            ot.setStyle(TableStyle([
                                ("LEFTPADDING",   (0,0),(-1,-1), 18),
                                ("TOPPADDING",    (0,0),(-1,-1), 1.5),
                                ("BOTTOMPADDING", (0,0),(-1,-1), 1.5),
                            ]))
                            items.append(ot)

            items.append(Spacer(1, 4 * mm))
            story.append(KeepTogether(items))

    # Marking Scheme
    story.append(PageBreak())
    story.append(Spacer(1, 2 * mm))
    story.append(HRFlowable(width=W, thickness=1.5, color=black))
    story.append(Paragraph("MARKING SCHEME", styles["ms_head"]))
    story.append(Paragraph(
        f"{_sanitize(data.get('subject',''))}  |  "
        f"{data.get('class_level','')}  |  Total Marks: {total_marks}",
        styles["meta"]
    ))
    story.append(Spacer(1, 2 * mm))
    story.append(HRFlowable(width=W, thickness=1.5, color=black))
    story.append(Spacer(1, 4 * mm))

    for entry in data.get("marking_scheme", []):
        qn        = entry.get("question_number", "?")
        marks     = entry.get("marks", "?")
        q_sum     = _prep(entry.get("question_summary", ""))
        ans_key   = _prep(entry.get("answer_key", ""))
        breakdown = [_prep(s) for s in entry.get("mark_breakdown", [])]
        q_type    = entry.get("type", "")

        story.append(Paragraph(
            f"<b>Q{qn}.</b> [{marks} {'Mark' if marks == 1 else 'Marks'}]"
            f" <i>({q_type})</i>  --  {q_sum}",
            styles["ms_q"]
        ))
        if ans_key:
            story.append(Paragraph(f"<b>Answer / Key:</b> {ans_key}", styles["ms_answer"]))
        if breakdown:
            story.append(Paragraph("<b>Mark Breakdown:</b>", styles["ms_answer"]))
            for step in breakdown:
                story.append(Paragraph(f"* {step}", styles["ms_step"]))
        story.append(HRFlowable(width=W, thickness=0.3, color=black))
        story.append(Spacer(1, 2 * mm))

    return story


# ─────────────────────────────────────────────
# PDF builder (two-pass for Page X of Y)
# ─────────────────────────────────────────────
def build_pdf(data: dict) -> bytes:
    buffer     = io.BytesIO()
    styles     = _get_styles()
    subject    = data.get("subject", "General")
    class_lv   = data.get("class_level", "")
    topic_str  = data.get("topic", "")
    month_year = datetime.now().strftime("%B %Y")
    total_pages: dict = {"n": 0}

    cb_first, cb_later = make_callbacks(subject, class_lv, topic_str, month_year, total_pages)

    def make_doc(buf):
        return SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=MARGIN_L, rightMargin=MARGIN_R,
            topMargin=MARGIN_T,  bottomMargin=MARGIN_B,
            title=data.get("title", "Assignment"), author=SCHOOL_NAME,
        )

    # Pass 1 – count pages
    count_doc = make_doc(io.BytesIO())
    count_doc.build(_build_story(data, styles),
                    onFirstPage=cb_first, onLaterPages=cb_later)
    total_pages["n"] = count_doc.page

    # Pass 2 – real render
    make_doc(buffer).build(_build_story(data, styles),
                           onFirstPage=cb_first, onLaterPages=cb_later)
    buffer.seek(0)
    return buffer.read()


# ─────────────────────────────────────────────
# DOCX builder
# ─────────────────────────────────────────────
def build_docx(data: dict) -> bytes:
    if not DOCX_OK:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    doc = DocxDocument()
    for sec in doc.sections:
        sec.left_margin  = Cm(2)
        sec.right_margin = Cm(2)
        sec.top_margin   = Cm(2.5)
        sec.bottom_margin= Cm(2)

    # Header
    hdr = doc.sections[0].header
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(SCHOOL_NAME)
    r.bold = True; r.font.size = Pt(13)
    hp.add_run(f"\n{data.get('subject','')} Assignment  |  "
               f"{data.get('class_level','')}  |  "
               f"{datetime.now().strftime('%B %Y')}\n"
               f"Topic: {data.get('topic','')}").font.size = Pt(9)

    # Title
    subject     = data.get("subject", "Physics")
    topic_clean = _clean_topic(data.get("title", ""), data.get("topic", ""))

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(f"{subject} Assignment")
    r_sub.bold = True; r_sub.font.size = Pt(12)

    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_top = p_top.add_run(topic_clean)
    r_top.bold = True; r_top.font.size = Pt(14)
    doc.add_paragraph()

    total_marks = data.get("total_marks", "-")

    for section in data.get("sections", []):
        doc.add_heading(section.get("section_label", ""), 2)
        ip = doc.add_paragraph(section.get("instructions", ""))
        if ip.runs:
            ip.runs[0].italic = True

        for q in section.get("questions", []):
            q_num  = q.get("number", "")
            q_text = q.get("question", "")
            marks  = q.get("marks", 1)
            q_type = q.get("type", "")

            p = doc.add_paragraph()
            run = p.add_run(
                f"Q{q_num}. {q_text}  "
                f"[{marks} {'Mark' if marks == 1 else 'Marks'}]"
            )
            run.bold = False

            if q_type == "MCQ":
                for opt in q.get("options", []):
                    op = doc.add_paragraph(opt)
                    op.paragraph_format.left_indent  = Cm(1)
                    op.paragraph_format.space_before = Pt(2)
                    op.paragraph_format.space_after  = Pt(2)

            doc.add_paragraph()

    # Marking scheme
    doc.add_page_break()
    mh = doc.add_heading("MARKING SCHEME", 0)
    mh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mi = doc.add_paragraph(
        f"{data.get('subject','')}  |  "
        f"{data.get('class_level','')}  |  Total Marks: {total_marks}"
    )
    mi.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for entry in data.get("marking_scheme", []):
        qn    = entry.get("question_number", "?")
        marks = entry.get("marks", "?")
        q_sum = entry.get("question_summary", "")
        ans   = entry.get("answer_key", "")
        bd    = entry.get("mark_breakdown", [])
        q_type= entry.get("type", "")

        p = doc.add_paragraph()
        run = p.add_run(
            f"Q{qn}. [{marks} {'Mark' if marks == 1 else 'Marks'}] "
            f"({q_type})  --  {q_sum}"
        )
        run.bold = True

        if ans:
            ap = doc.add_paragraph(f"Answer / Key: {ans}")
            ap.paragraph_format.left_indent = Cm(0.5)

        if bd:
            bp = doc.add_paragraph("Mark Breakdown:")
            if bp.runs: bp.runs[0].bold = True
            bp.paragraph_format.left_indent = Cm(0.5)
            for step in bd:
                sp = doc.add_paragraph(f"* {step}")
                sp.paragraph_format.left_indent = Cm(1)

        doc.add_paragraph("-" * 60)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# Error classifier
# ─────────────────────────────────────────────
def classify_error(e: Exception):
    s   = str(e).lower()
    raw = str(e)

    if "gemini_api_key_missing" in s:
        return ("INVALID_KEY",
                "Gemini API key is not set. Please add GEMINI_API_KEY to your .env file.",
                "Get a free key at https://aistudio.google.com/app/apikey and paste it into .env.", 401)

    if any(k in s for k in ("429","resource_exhausted","quota","rate limit","too many requests","ratelimitexceeded")):
        return ("RATE_LIMIT",
                "You've hit the Gemini free-tier rate limit. This is a temporary restriction.",
                "Wait 60-120 seconds and try again. If it keeps happening, upgrade at https://aistudio.google.com.", 429)

    if any(k in s for k in ("api_key","invalid api key","unauthenticated","permission_denied","403","401")):
        return ("INVALID_KEY",
                "Invalid or missing Gemini API key.",
                "Copy a fresh key from https://aistudio.google.com/app/apikey and update your .env file.", 401)

    if isinstance(e, json.JSONDecodeError) or "json" in s:
        return ("JSON_PARSE",
                "Gemini returned a response that couldn't be parsed.",
                "Try reducing the number of questions or simplifying the topic, then retry.", 502)

    if any(k in s for k in ("timeout","timed out","connection","network","connectionerror")):
        return ("TIMEOUT",
                "The request to Gemini timed out or the network is unavailable.",
                "Check your internet connection and try again.", 504)

    if any(k in s for k in ("reportlab","flowable","platypus","canvas")):
        return ("PDF_ERROR",
                "There was a problem rendering the PDF.",
                "Try again. If it persists, reduce the question count or shorten the topic.", 500)

    if any(k in s for k in ("503","service unavailable","overloaded","server_error","internal")):
        return ("SERVICE_DOWN",
                "The Gemini API is currently overloaded or temporarily unavailable.",
                "Wait a few minutes and try again. Check https://status.cloud.google.com for outages.", 503)

    return ("UNKNOWN", f"An unexpected error occurred: {raw}",
            "Check the server console for the full traceback.", 500)


# ─────────────────────────────────────────────
# Flask routes & Universal Catch-All Handler
# ─────────────────────────────────────────────
def _handle_generate():
    try:
        body = request.get_json(force=True)
        class_level         = body.get("class_level", "Class 10")
        topic               = body.get("topic", "").strip()
        marks_2             = int(body.get("marks_2", body.get("marks_1", 0)))
        marks_3             = int(body.get("marks_3", 0))
        marks_5             = int(body.get("marks_5", 0))
        mcq                 = int(body.get("mcq", 0))
        difficulty_level    = body.get("difficulty_level", "standard")
        subtopics           = body.get("subtopics", "").strip()
        exclusions          = body.get("exclusions", "").strip()
        custom_instructions = body.get("custom_instructions", "").strip()

        if not topic:
            return jsonify({"error":"Topic is required.","error_code":"MISSING_TOPIC",
                            "hint":"Type the chapter or topic name in the Main Topic field."}), 400

        if marks_2 + marks_3 + marks_5 + mcq == 0:
            return jsonify({"error":"Please add at least one question using the sliders.",
                            "error_code":"NO_QUESTIONS",
                            "hint":"Move any slider above 0 before generating."}), 400

        prompt = build_prompt(class_level, topic, marks_2, marks_3, marks_5, mcq,
                              difficulty_level, subtopics, exclusions, custom_instructions)
        data = call_gemini(prompt)

        pdf_bytes  = build_pdf(data)
        pdf_b64    = base64.b64encode(pdf_bytes).decode("utf-8")

        docx_b64 = ""
        if DOCX_OK:
            try:
                docx_bytes = build_docx(data)
                docx_b64   = base64.b64encode(docx_bytes).decode("utf-8")
            except Exception:
                pass  # DOCX failure is non-fatal; PDF still delivered

        return jsonify({
            "success":     True,
            "pdf_base64":  pdf_b64,
            "docx_base64": docx_b64,
            "title":       data.get("title", "Assignment"),
            "total_marks": data.get("total_marks", "-"),
            "subject":     data.get("subject", "-"),
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        code, msg, hint, status = classify_error(e)
        return jsonify({"error": msg, "error_code": code, "hint": hint}), status


@app.route("/", defaults={"path": ""}, methods=["GET", "POST"])
@app.route("/<path:path>", methods=["GET", "POST"])
def catch_all(path):
    if "generate" in path and request.method == "POST":
        return _handle_generate()
    if "favicon" in path:
        return "", 204
    return render_template("index.html")


@app.route("/health")
def health():
    return jsonify({"status": "ok", "api_key_set": bool(GEMINI_API_KEY)})


if __name__ == "__main__":
    app.run(debug=True, port=5000)
